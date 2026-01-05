#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Forguncy Insight - Forguncyプロジェクト解析・仕様書自動生成ツール
Python版 - 完全版（ライセンス管理・Excel出力・差分比較対応）
"""

# =============================================================================
# バージョン情報
# =============================================================================
APP_VERSION = "1.0.0"
SUPPORTED_FORGUNCY_VERSIONS = ["9.x"]  # 対応Forguncyバージョン
FORGUNCY_VERSION_TESTED = "9.0"  # テスト済みバージョン
VERSION_INFO = f"v{APP_VERSION} (Forguncy {', '.join(SUPPORTED_FORGUNCY_VERSIONS)} 対応)"

import hashlib
import json
import os
import re
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from tkinter import Tk, Frame, Label, Button, Entry, StringVar, filedialog, messagebox, ttk, Toplevel, Text, Scrollbar
from tkinter import LEFT, RIGHT, BOTH, END, X, Y, W, E, N, S, VERTICAL, HORIZONTAL, WORD
from typing import Any, Optional, Dict, List, Tuple
import webbrowser

# ドラッグ＆ドロップサポート
try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False

# python-docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# openpyxl for Excel
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


# =============================================================================
# ライセンス管理（新形式: PPPP-PLAN-YYMM-HASH-SIG1-SIG2）
# =============================================================================
import hmac
import base64

# ライセンスキー正規表現
# 形式: PPPP-PLAN-YYMM-HASH-SIG1-SIG2
LICENSE_KEY_PATTERN = re.compile(
    r'^(INSS|INSP|INPY|FGIN)-(TRIAL|STD|PRO)-(\d{4})-([A-Z0-9]{4})-([A-Z0-9]{4})-([A-Z0-9]{4})$'
)

PRODUCT_CODE = 'FGIN'  # Forguncy Insight
PRODUCT_NAME = 'Forguncy Insight'
PRICE_STANDARD = '¥49,800/年'

# 購入・トライアルURL
PURCHASE_URL = 'https://because-zero.com/forguncy-insight/purchase'
TRIAL_URL = 'https://because-zero.com/forguncy-insight/trial'

# 期限警告の日数
EXPIRY_WARNING_DAYS = 30

# トライアル期間（日）
TRIAL_DAYS = 14

# 署名用シークレットキー
_SECRET_KEY = os.environ.get(
    "INSIGHT_LICENSE_SECRET",
    b"insight-series-license-secret-2026"
)
if isinstance(_SECRET_KEY, str):
    _SECRET_KEY = _SECRET_KEY.encode()

# 機能制限
FEATURE_LIMITS = {
    'FREE': {
        'max_tables': 5,
        'max_pages': 10,
        'max_server_commands': 3,
        'max_workflows': 1,
        'word_export': False,
        'excel_export': False,
        'diff_compare': False,
        'commercial_use': False,
    },
    'STD': {
        'max_tables': float('inf'),
        'max_pages': float('inf'),
        'max_server_commands': float('inf'),
        'max_workflows': float('inf'),
        'word_export': True,
        'excel_export': True,
        'diff_compare': True,
        'commercial_use': True,
    },
}

TIER_NAMES = {
    'TRIAL': 'トライアル',
    'STD': 'Standard',
    'PRO': 'Pro',
}


def _generate_email_hash(email: str) -> str:
    """メールアドレスから4文字のハッシュを生成"""
    h = hashlib.sha256(email.lower().strip().encode()).digest()
    return base64.b32encode(h)[:4].decode().upper()


def _generate_signature(data: str) -> str:
    """署名を生成（8文字）"""
    sig = hmac.new(_SECRET_KEY, data.encode(), hashlib.sha256).digest()
    encoded = base64.b32encode(sig)[:8].decode().upper()
    return encoded


def _verify_signature(data: str, signature: str) -> bool:
    """署名を検証"""
    expected = _generate_signature(data)
    return hmac.compare_digest(expected, signature)


def validate_license_key(license_key: str, email: str) -> Dict:
    """
    ライセンスキーを検証（新形式）
    形式: PPPP-PLAN-YYMM-HASH-SIG1-SIG2
    """
    if not license_key:
        return {'is_valid': False, 'error': 'ライセンスキーが必要です', 'tier': None}

    if not email:
        return {'is_valid': False, 'error': 'メールアドレスが必要です', 'tier': None}

    normalized = license_key.strip().upper()
    email = email.strip().lower()
    match = LICENSE_KEY_PATTERN.match(normalized)

    if not match:
        return {'is_valid': False, 'error': '無効なライセンスキー形式です', 'tier': None}

    product_code, tier, yymm, email_hash, sig1, sig2 = match.groups()
    signature = sig1 + sig2

    # 署名検証
    sign_data = f"{product_code}-{tier}-{yymm}-{email_hash}"
    if not _verify_signature(sign_data, signature):
        return {'is_valid': False, 'error': 'ライセンスキーが無効です', 'tier': None}

    # メールハッシュ照合
    expected_hash = _generate_email_hash(email)
    if email_hash != expected_hash:
        return {'is_valid': False, 'error': 'メールアドレスが一致しません', 'tier': None}

    # 有効期限チェック
    try:
        year = 2000 + int(yymm[:2])
        month = int(yymm[2:])
        # 月末日を有効期限とする
        if month == 12:
            expires_at = datetime(year + 1, 1, 1) - timedelta(days=1)
        else:
            expires_at = datetime(year, month + 1, 1) - timedelta(days=1)
        expires_at = expires_at.replace(hour=23, minute=59, second=59)
    except ValueError:
        return {'is_valid': False, 'error': '無効な有効期限です', 'tier': None}

    if datetime.now() > expires_at:
        return {
            'is_valid': False,
            'error': 'ライセンスの有効期限が切れています',
            'tier': tier,
            'expires_at': expires_at
        }

    # 製品コードチェック
    if product_code != PRODUCT_CODE:
        return {'is_valid': False, 'error': f'このライセンスは {product_code} 用です', 'tier': None}

    return {
        'is_valid': True,
        'tier': tier,
        'product': product_code,
        'expires_at': expires_at,
        'error': None
    }


def get_feature_limits(tier: Optional[str]) -> Dict:
    """ティアに応じた機能制限を取得"""
    if tier in ['TRIAL', 'STD', 'PRO']:
        return FEATURE_LIMITS['STD']
    return FEATURE_LIMITS['FREE']


class LicenseManager:
    """ライセンス管理クラス（新形式対応）"""

    def __init__(self):
        self._config_dir = self._get_config_dir()
        self._config_file = self._config_dir / 'license.dat'
        self.license_key = None
        self.email = None
        self.expires_at = None
        self.activated_at = None
        self.tier = None
        self.limits = get_feature_limits(None)
        self.load()

    def _get_config_dir(self) -> Path:
        """設定ディレクトリを取得"""
        if os.name == 'nt':  # Windows
            base = Path(os.environ.get('APPDATA', Path.home()))
        else:
            base = Path.home() / '.config'
        return base / 'HarmonicInsight' / 'ForguncyInsight'

    def load(self):
        """保存されたライセンス情報を読み込み"""
        if not self._config_file.exists():
            return

        try:
            with open(self._config_file, 'r', encoding='utf-8') as f:
                encoded = f.read()
            content = base64.b64decode(encoded).decode()
            data = json.loads(content)

            self.license_key = data.get('key')
            self.email = data.get('email')
            self.tier = data.get('plan')
            expires_str = data.get('expires')
            if expires_str:
                self.expires_at = datetime.strptime(expires_str, '%Y-%m-%d')
                self.expires_at = self.expires_at.replace(hour=23, minute=59, second=59)

            # 期限チェック
            if self.expires_at and datetime.now() > self.expires_at:
                self.tier = None
                self.limits = get_feature_limits(None)
            else:
                self.limits = get_feature_limits(self.tier)
        except Exception:
            pass

    def save(self):
        """ライセンス情報を保存"""
        self._config_dir.mkdir(parents=True, exist_ok=True)

        data = {
            'key': self.license_key,
            'email': self.email,
            'plan': self.tier,
            'productCode': PRODUCT_CODE,
            'product': PRODUCT_NAME,
            'expires': self.expires_at.strftime('%Y-%m-%d') if self.expires_at else None,
            'verifiedAt': datetime.now().isoformat()
        }

        content = json.dumps(data, ensure_ascii=False)
        encoded = base64.b64encode(content.encode()).decode()

        with open(self._config_file, 'w', encoding='utf-8') as f:
            f.write(encoded)

    def activate(self, email: str, license_key: str) -> Dict:
        """ライセンスをアクティベート"""
        # メールアドレスの検証
        if not email or '@' not in email:
            return {'is_valid': False, 'error': '有効なメールアドレスを入力してください', 'tier': None}

        # 新形式でバリデーション
        info = validate_license_key(license_key, email)

        if info['is_valid']:
            self.license_key = license_key.strip().upper()
            self.email = email.strip().lower()
            self.expires_at = info.get('expires_at')
            self.activated_at = datetime.now()
            self.tier = info.get('tier')
            self.limits = get_feature_limits(self.tier)
            self.save()

        return info

    def clear(self):
        """ライセンスをクリア"""
        self.license_key = None
        self.email = None
        self.expires_at = None
        self.activated_at = None
        self.tier = None
        self.limits = get_feature_limits(None)
        if self._config_file.exists():
            self._config_file.unlink()

    @property
    def is_activated(self) -> bool:
        return self.tier is not None

    @property
    def tier_name(self) -> str:
        if self.tier:
            return TIER_NAMES.get(self.tier, self.tier)
        return 'Free'

    @property
    def days_until_expiry(self) -> Optional[int]:
        """有効期限までの日数を返す"""
        if not self.expires_at:
            return None
        delta = self.expires_at - datetime.now()
        return delta.days

    @property
    def is_expiring_soon(self) -> bool:
        """期限が近づいているかチェック（30日前から警告）"""
        days = self.days_until_expiry
        if days is None:
            return False
        return 0 < days <= EXPIRY_WARNING_DAYS

    @property
    def expiry_warning_message(self) -> Optional[str]:
        """期限警告メッセージを返す"""
        if not self.is_expiring_soon:
            return None
        days = self.days_until_expiry
        return f"ライセンスの有効期限まであと{days}日です。更新をお忘れなく！"


# =============================================================================
# データクラス定義
# =============================================================================

@dataclass
class ColumnInfo:
    name: str
    type: str
    required: bool = False
    unique: bool = False
    default_value: Optional[str] = None
    description: Optional[str] = None


@dataclass
class RelationInfo:
    target_table: str
    source_column: str
    target_column: str
    relation_type: str = "OneToMany"


@dataclass
class CommandInfo:
    type: str
    description: str
    details: dict = field(default_factory=dict)
    sub_commands: list = field(default_factory=list)


@dataclass
class StateInfo:
    name: str
    is_initial: bool = False
    is_final: bool = False


@dataclass
class AssigneeInfo:
    type: str
    value: str


@dataclass
class ConditionInfo:
    type: str
    field: Optional[str] = None
    operator: Optional[str] = None
    value: Optional[str] = None
    expression: Optional[str] = None


@dataclass
class TransitionInfo:
    from_state: str
    to_state: str
    action: str
    conditions: list = field(default_factory=list)
    assignees: list = field(default_factory=list)
    commands: list = field(default_factory=list)


@dataclass
class WorkflowInfo:
    table_name: str
    states: list = field(default_factory=list)
    transitions: list = field(default_factory=list)


@dataclass
class TableInfo:
    name: str
    folder: str = ""
    columns: list = field(default_factory=list)
    relations: list = field(default_factory=list)
    workflow: Optional[WorkflowInfo] = None


@dataclass
class FormulaInfo:
    cell: str
    formula: str


@dataclass
class ButtonInfo:
    name: str
    cell: str = ""
    commands: list = field(default_factory=list)


@dataclass
class CellCommandInfo:
    cell: str
    event: str
    commands: list = field(default_factory=list)


@dataclass
class PageInfo:
    name: str
    page_type: str
    path: str
    buttons: list = field(default_factory=list)
    formulas: list = field(default_factory=list)
    cell_commands: list = field(default_factory=list)


@dataclass
class ParameterInfo:
    name: str
    type: str
    required: bool = False
    default_value: Optional[str] = None


@dataclass
class ServerCommandInfo:
    name: str
    folder: str = ""
    path: str = ""
    commands: list = field(default_factory=list)
    raw_commands: list = field(default_factory=list)
    parameters: list = field(default_factory=list)


@dataclass
class AnalysisSummary:
    table_count: int = 0
    page_count: int = 0
    workflow_count: int = 0
    server_command_count: int = 0
    total_columns: int = 0
    total_relations: int = 0


@dataclass
class AnalysisResult:
    project_name: str
    tables: list = field(default_factory=list)
    pages: list = field(default_factory=list)
    workflows: list = field(default_factory=list)
    server_commands: list = field(default_factory=list)
    summary: AnalysisSummary = field(default_factory=AnalysisSummary)


# =============================================================================
# 解析ロジック
# =============================================================================

def extract_json(content: str) -> dict:
    """Forguncy特殊形式のJSONを抽出"""
    cleaned = content.lstrip('\ufeff')
    start = cleaned.find('{')
    if start == -1:
        raise ValueError("No JSON object found")

    brace_count = 0
    in_string = False
    escape = False

    for i in range(start, len(cleaned)):
        char = cleaned[i]
        if escape:
            escape = False
            continue
        if char == '\\' and in_string:
            escape = True
            continue
        if char == '"':
            in_string = not in_string
            continue
        if not in_string:
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    return json.loads(cleaned[start:i+1])

    raise ValueError("Invalid JSON structure")


def extract_column_type(column_type: Optional[str]) -> str:
    if not column_type:
        return "Text"
    parts = column_type.split(',')[0].split('.')
    return parts[-1] if parts else "Text"


def extract_command_type_name(type_string: str) -> str:
    if not type_string:
        return "Unknown"
    parts = type_string.split(',')[0].split('.')
    return parts[-1] if parts else "Unknown"


def format_condition(condition: Optional[dict]) -> str:
    if not condition:
        return "(条件なし)"
    if condition.get('Expression'):
        return str(condition['Expression'])
    left = condition.get('LeftOperand', '')
    op = condition.get('Operator', '==')
    right = condition.get('RightOperand', '')
    return f"{left} {op} {right}"


def parse_conditions(condition: Optional[dict]) -> list:
    if not condition:
        return []
    conditions = []
    cond_type = condition.get('$type', '')
    if 'ExpressionCondition' in cond_type:
        conditions.append(ConditionInfo(type='expression', expression=condition.get('Expression')))
    elif 'CompareCondition' in cond_type:
        conditions.append(ConditionInfo(
            type='compare',
            field=str(condition.get('LeftOperand', '')),
            operator=condition.get('Operator', '=='),
            value=str(condition.get('RightOperand', ''))
        ))
    return conditions


def parse_assignees(assignees: list) -> list:
    result = []
    for a in assignees:
        a_type = a.get('$type', '')
        if 'UserAssignee' in a_type:
            result.append(AssigneeInfo(type='user', value=a.get('UserName', '')))
        elif 'RoleAssignee' in a_type:
            result.append(AssigneeInfo(type='role', value=a.get('RoleName', '')))
        elif 'FieldAssignee' in a_type:
            result.append(AssigneeInfo(type='field', value=a.get('FieldName', '')))
        elif 'CreatorAssignee' in a_type:
            result.append(AssigneeInfo(type='creator', value='作成者'))
        elif 'PreviousAssignee' in a_type:
            result.append(AssigneeInfo(type='previousAssignee', value='前の担当者'))
        else:
            result.append(AssigneeInfo(type='user', value='不明'))
    return result


def generate_command_description(cmd: dict, type_name: str) -> str:
    descriptions = {
        'ExecuteSqlCommand': 'SQL実行',
        'UpdateTableDataCommand': f"テーブル更新: {cmd.get('TableName', '')}",
        'InsertTableDataCommand': f"テーブル挿入: {cmd.get('TableName', '')}",
        'DeleteTableDataCommand': f"テーブル削除: {cmd.get('TableName', '')}",
        'SendEmailCommand': 'メール送信',
        'ConditionCommand': '条件分岐',
        'LoopCommand': 'ループ処理',
        'SetCellValueCommand': 'セル値設定',
        'NavigateCommand': 'ページ遷移',
        'CallServerCommandCommand': 'サーバーコマンド呼出',
    }
    return descriptions.get(type_name, type_name)


def parse_command(cmd: dict) -> CommandInfo:
    cmd_type = cmd.get('$type', '')
    type_name = extract_command_type_name(cmd_type)
    command = CommandInfo(type=type_name, description=generate_command_description(cmd, type_name), details={})

    if type_name == 'ConditionCommand' or 'ConditionCommand' in cmd_type:
        command.description = f"IF {format_condition(cmd.get('Condition'))}"
        command.sub_commands = parse_commands(cmd.get('TrueCommands', [])) + parse_commands(cmd.get('FalseCommands', []))

    if 'ExecuteSqlCommand' in cmd_type:
        sql = cmd.get('SqlStatement', '')
        command.details = {'sql': sql}
        command.description = f"SQL実行: {sql[:100]}..."

    if 'UpdateTableDataCommand' in cmd_type:
        command.details = {'table': cmd.get('TableName'), 'mappings': cmd.get('ColumnMappings')}
        command.description = f"テーブル更新: {cmd.get('TableName')}"

    if 'SendEmailCommand' in cmd_type:
        command.details = {'to': cmd.get('EmailTo'), 'subject': cmd.get('EmailSubject')}
        command.description = f"メール送信: {cmd.get('EmailSubject', '(件名なし)')}"

    return command


def parse_commands(commands: list) -> list:
    return [parse_command(cmd) for cmd in commands]


def parse_workflow(table_name: str, wf_data: dict) -> WorkflowInfo:
    states = [StateInfo(name=s.get('Name', ''), is_initial=s.get('IsInitialState', False), is_final=s.get('IsFinalState', False)) for s in wf_data.get('States', [])]
    transitions = [TransitionInfo(
        from_state=t.get('SourceStateName', ''),
        to_state=t.get('TargetStateName', ''),
        action=t.get('ActionName', ''),
        conditions=parse_conditions(t.get('Condition')),
        assignees=parse_assignees(t.get('Assignees', [])),
        commands=parse_commands(t.get('Commands', []))
    ) for t in wf_data.get('Transitions', [])]
    return WorkflowInfo(table_name=table_name, states=states, transitions=transitions)


def flatten_commands_to_text(commands: list, depth: int = 0) -> list:
    lines = []
    indent = '  ' * depth
    for cmd in commands:
        cmd_type = cmd.get('$type', '')
        type_name = extract_command_type_name(cmd_type)
        if 'ConditionCommand' in cmd_type:
            lines.append(f"{indent}IF {format_condition(cmd.get('Condition'))} THEN")
            if cmd.get('TrueCommands'):
                lines.extend(flatten_commands_to_text(cmd['TrueCommands'], depth + 1))
            if cmd.get('FalseCommands'):
                lines.append(f"{indent}ELSE")
                lines.extend(flatten_commands_to_text(cmd['FalseCommands'], depth + 1))
            lines.append(f"{indent}END IF")
        elif 'LoopCommand' in cmd_type:
            lines.append(f"{indent}LOOP")
            if cmd.get('Commands'):
                lines.extend(flatten_commands_to_text(cmd['Commands'], depth + 1))
            lines.append(f"{indent}END LOOP")
        elif 'ExecuteSqlCommand' in cmd_type:
            lines.append(f"{indent}EXECUTE SQL:")
            for line in cmd.get('SqlStatement', '').split('\n'):
                lines.append(f"{indent}  {line}")
        elif 'UpdateTableDataCommand' in cmd_type:
            lines.append(f"{indent}UPDATE TABLE: {cmd.get('TableName')}")
        elif 'InsertTableDataCommand' in cmd_type:
            lines.append(f"{indent}INSERT INTO TABLE: {cmd.get('TableName')}")
        elif 'DeleteTableDataCommand' in cmd_type:
            lines.append(f"{indent}DELETE FROM TABLE: {cmd.get('TableName')}")
        elif 'SendEmailCommand' in cmd_type:
            lines.append(f"{indent}SEND EMAIL TO: {cmd.get('EmailTo')}")
            lines.append(f"{indent}  SUBJECT: {cmd.get('EmailSubject')}")
        elif 'CallServerCommandCommand' in cmd_type:
            lines.append(f"{indent}CALL SERVER COMMAND: {cmd.get('ServerCommandName', '(不明)')}")
        else:
            lines.append(f"{indent}{type_name}")
    return lines


def infer_parameter_type(validation_info: Optional[dict]) -> str:
    if not validation_info:
        return 'Text'
    number_type = validation_info.get('NumberType')
    return {4: 'DateTime', 1: 'Integer', 2: 'Decimal'}.get(number_type, 'Text')


def analyze_project(file_path: str, progress_callback=None, limits=None) -> AnalysisResult:
    """Forguncyプロジェクトを解析"""
    def send_progress(pct, msg):
        if progress_callback:
            progress_callback(pct, msg)

    project_name = Path(file_path).stem
    limits = limits or FEATURE_LIMITS['FREE']

    with zipfile.ZipFile(file_path, 'r') as zf:
        entries = zf.namelist()

        send_progress(15, 'テーブル定義を解析しています...')
        tables = analyze_tables(zf, entries, int(limits.get('max_tables', 5)))

        send_progress(25, 'ページ定義を解析しています...')
        pages = analyze_pages(zf, entries, int(limits.get('max_pages', 10)))

        send_progress(35, 'ワークフローを解析しています...')
        max_wf = int(limits.get('max_workflows', 1))
        workflows = [t.workflow for t in tables if t.workflow][:max_wf]

        send_progress(45, 'サーバーコマンドを解析しています...')
        server_commands = analyze_server_commands(zf, entries, int(limits.get('max_server_commands', 3)))

    summary = AnalysisSummary(
        table_count=len(tables),
        page_count=len(pages),
        workflow_count=len(workflows),
        server_command_count=len(server_commands),
        total_columns=sum(len(t.columns) for t in tables),
        total_relations=sum(len(t.relations) for t in tables)
    )

    return AnalysisResult(project_name=project_name, tables=tables, pages=pages, workflows=workflows, server_commands=server_commands, summary=summary)


def analyze_tables(zf: zipfile.ZipFile, entries: list, max_count: int = 999) -> list:
    tables = []
    table_entries = [e for e in entries if e.startswith('Tables/') and e.endswith('.json')][:max_count]

    for entry in table_entries:
        try:
            content = zf.read(entry).decode('utf-8')
            data = extract_json(content)
            path_parts = entry.split('/')
            folder = path_parts[1] if len(path_parts) > 2 else ''

            columns = [ColumnInfo(
                name=col.get('Name', ''),
                type=extract_column_type(col.get('ColumnType')),
                required=col.get('Required', False),
                unique=col.get('Unique', False),
                default_value=str(col['DefaultValue']) if col.get('DefaultValue') is not None else None,
                description=col.get('Description')
            ) for col in data.get('Columns', [])]

            relations = [RelationInfo(
                target_table=rel.get('TargetTableName', ''),
                source_column=rel.get('SourceColumnName', ''),
                target_column=rel.get('TargetColumnName', ''),
                relation_type=rel.get('RelationType', 'OneToMany')
            ) for rel in data.get('Relations', [])]

            table = TableInfo(name=data.get('Name', Path(entry).stem), folder=folder, columns=columns, relations=relations)

            if data.get('BindingRelatedWorkflow'):
                table.workflow = parse_workflow(table.name, data['BindingRelatedWorkflow'])

            tables.append(table)
        except Exception as e:
            print(f"Error parsing table {entry}: {e}")

    return tables


def analyze_pages(zf: zipfile.ZipFile, entries: list, max_count: int = 999) -> list:
    pages = []

    for entry in [e for e in entries if e.startswith('Pages/') and e.endswith('.json')][:max_count]:
        try:
            content = zf.read(entry).decode('utf-8')
            data = extract_json(content)
            elements = extract_page_elements(data)
            pages.append(PageInfo(name=data.get('Name', Path(entry).stem), page_type='page', path=entry, **elements))
        except Exception as e:
            print(f"Error parsing page {entry}: {e}")

    for entry in [e for e in entries if e.startswith('MasterPages/') and e.endswith('.json')]:
        try:
            content = zf.read(entry).decode('utf-8')
            data = extract_json(content)
            elements = extract_page_elements(data)
            pages.append(PageInfo(name=data.get('Name', Path(entry).stem), page_type='masterPage', path=entry, **elements))
        except Exception as e:
            print(f"Error parsing master page {entry}: {e}")

    return pages


def extract_page_elements(data: dict) -> dict:
    buttons, formulas, cell_commands = [], [], []
    attach_infos = data.get('AttachInfos', {})

    for cell_address, cell_data in attach_infos.items():
        cell_type = cell_data.get('CellType', {})
        if cell_data.get('Formula'):
            formulas.append(FormulaInfo(cell=cell_address, formula=str(cell_data['Formula'])))

        if cell_type:
            type_str = cell_type.get('$type', '')
            if 'MenuCellType' in type_str or 'ForguncyMenuCellType' in type_str:
                extract_menu_items(cell_type.get('Items', []), buttons, cell_address)
            if 'ButtonCellType' in type_str:
                text = cell_type.get('Text') or cell_type.get('Content') or 'ボタン'
                command_list = cell_type.get('CommandList', [])
                if command_list:
                    buttons.append(ButtonInfo(name=text, cell=cell_address, commands=parse_commands(command_list)))
            command_list = cell_type.get('CommandList', [])
            if command_list and 'ButtonCellType' not in type_str:
                cell_commands.append(CellCommandInfo(cell=cell_address, event='Click', commands=parse_commands(command_list)))

    return {'buttons': buttons, 'formulas': formulas, 'cell_commands': cell_commands}


def extract_menu_items(items: list, buttons: list, base_cell: str):
    for item in items:
        if item.get('CommandList'):
            buttons.append(ButtonInfo(name=f"メニュー: {item.get('Text', '(名称なし)')}", cell=base_cell, commands=parse_commands(item['CommandList'])))
        if item.get('SubItems'):
            extract_menu_items(item['SubItems'], buttons, base_cell)


def analyze_server_commands(zf: zipfile.ZipFile, entries: list, max_count: int = 999) -> list:
    server_commands = []
    cmd_entries = [e for e in entries if e.startswith('ServerCommands/') and e.endswith('.json')][:max_count]

    for entry in cmd_entries:
        try:
            content = zf.read(entry).decode('utf-8')
            data = extract_json(content)
            path_parts = entry.split('/')
            folder = path_parts[1] if len(path_parts) > 2 else ''

            raw_commands = parse_commands(data.get('Commands', []))
            commands = flatten_commands_to_text(data.get('Commands', []))

            parameters = []
            triggers = data.get('Triggers', [])
            if triggers:
                for p in triggers[0].get('Parameters', []):
                    parameters.append(ParameterInfo(name=p.get('Name', ''), type=infer_parameter_type(p.get('DataValidationInfo')), required=True))

            if not parameters and data.get('Parameters'):
                for p in data['Parameters']:
                    parameters.append(ParameterInfo(name=p.get('Name', ''), type=extract_column_type(p.get('Type')), required=p.get('Required', False), default_value=str(p['DefaultValue']) if p.get('DefaultValue') is not None else None))

            server_commands.append(ServerCommandInfo(name=data.get('Name', Path(entry).stem), folder=folder, path=entry, commands=commands, raw_commands=raw_commands, parameters=parameters))
        except Exception as e:
            print(f"Error parsing server command {entry}: {e}")

    return server_commands


# =============================================================================
# 差分比較
# =============================================================================

@dataclass
class DiffResult:
    added_tables: list = field(default_factory=list)
    removed_tables: list = field(default_factory=list)
    modified_tables: list = field(default_factory=list)
    added_pages: list = field(default_factory=list)
    removed_pages: list = field(default_factory=list)
    added_server_commands: list = field(default_factory=list)
    removed_server_commands: list = field(default_factory=list)
    modified_server_commands: list = field(default_factory=list)


def compare_projects(old_analysis: AnalysisResult, new_analysis: AnalysisResult) -> DiffResult:
    """2つのプロジェクトを比較"""
    diff = DiffResult()

    # テーブル比較
    old_tables = {t.name: t for t in old_analysis.tables}
    new_tables = {t.name: t for t in new_analysis.tables}

    diff.added_tables = [t for name, t in new_tables.items() if name not in old_tables]
    diff.removed_tables = [t for name, t in old_tables.items() if name not in new_tables]

    for name in set(old_tables.keys()) & set(new_tables.keys()):
        old_t, new_t = old_tables[name], new_tables[name]
        old_cols = {c.name: c for c in old_t.columns}
        new_cols = {c.name: c for c in new_t.columns}
        if old_cols != new_cols or len(old_t.relations) != len(new_t.relations):
            diff.modified_tables.append({'name': name, 'old': old_t, 'new': new_t})

    # ページ比較
    old_pages = {p.name for p in old_analysis.pages}
    new_pages = {p.name for p in new_analysis.pages}
    diff.added_pages = [p for p in new_analysis.pages if p.name not in old_pages]
    diff.removed_pages = [p for p in old_analysis.pages if p.name not in new_pages]

    # サーバーコマンド比較
    old_cmds = {c.name: c for c in old_analysis.server_commands}
    new_cmds = {c.name: c for c in new_analysis.server_commands}

    diff.added_server_commands = [c for name, c in new_cmds.items() if name not in old_cmds]
    diff.removed_server_commands = [c for name, c in old_cmds.items() if name not in new_cmds]

    for name in set(old_cmds.keys()) & set(new_cmds.keys()):
        if old_cmds[name].commands != new_cmds[name].commands:
            diff.modified_server_commands.append({'name': name, 'old': old_cmds[name], 'new': new_cmds[name]})

    return diff


# =============================================================================
# Word出力
# =============================================================================

def generate_spec_document(analysis: AnalysisResult, output_dir: str) -> str:
    """仕様書ドキュメントを生成"""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    style.font.size = Pt(11)

    # タイトル
    title = doc.add_heading(analysis.project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('システム仕様書').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'生成日: {datetime.now().strftime("%Y年%m月%d日")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_para = doc.add_paragraph(f'Forguncy Insight {VERSION_INFO} により自動生成')
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_para.runs[0].italic = True

    doc.add_page_break()

    # 目次
    doc.add_heading('目次', 1)
    doc.add_paragraph('1. システム概要')
    doc.add_paragraph('2. テーブル定義')
    doc.add_paragraph('3. 画面一覧')
    if analysis.workflows:
        doc.add_paragraph('4. ワークフロー定義')
    if analysis.server_commands:
        doc.add_paragraph('5. サーバーコマンド')

    doc.add_page_break()

    # 1. システム概要
    doc.add_heading('1. システム概要', 1)
    summary_table = doc.add_table(rows=8, cols=2)
    summary_table.style = 'Table Grid'
    for i, (label, value) in enumerate([
        ('プロジェクト名', analysis.project_name),
        ('テーブル数', f'{analysis.summary.table_count}件'),
        ('ページ数', f'{analysis.summary.page_count}件'),
        ('ワークフロー数', f'{analysis.summary.workflow_count}件'),
        ('サーバーコマンド数', f'{analysis.summary.server_command_count}件'),
        ('総カラム数', f'{analysis.summary.total_columns}件'),
        ('リレーション数', f'{analysis.summary.total_relations}件'),
        ('対応Forguncyバージョン', ', '.join(SUPPORTED_FORGUNCY_VERSIONS)),
    ]):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    doc.add_page_break()

    # 2. テーブル定義
    doc.add_heading('2. テーブル定義', 1)
    if analysis.tables:
        table_list = doc.add_table(rows=len(analysis.tables) + 1, cols=5)
        table_list.style = 'Table Grid'
        for j, h in enumerate(['No.', 'テーブル名', 'フォルダ', 'カラム数', 'リレーション']):
            table_list.rows[0].cells[j].text = h
        for i, t in enumerate(analysis.tables):
            table_list.rows[i + 1].cells[0].text = str(i + 1)
            table_list.rows[i + 1].cells[1].text = t.name
            table_list.rows[i + 1].cells[2].text = t.folder or '-'
            table_list.rows[i + 1].cells[3].text = str(len(t.columns))
            table_list.rows[i + 1].cells[4].text = str(len(t.relations))

    for table in analysis.tables:
        doc.add_heading(table.name, 3)
        if table.columns:
            col_table = doc.add_table(rows=len(table.columns) + 1, cols=5)
            col_table.style = 'Table Grid'
            for j, h in enumerate(['カラム名', 'データ型', '必須', 'ユニーク', 'デフォルト値']):
                col_table.rows[0].cells[j].text = h
            for i, c in enumerate(table.columns):
                col_table.rows[i + 1].cells[0].text = c.name
                col_table.rows[i + 1].cells[1].text = c.type
                col_table.rows[i + 1].cells[2].text = '○' if c.required else ''
                col_table.rows[i + 1].cells[3].text = '○' if c.unique else ''
                col_table.rows[i + 1].cells[4].text = c.default_value or ''

    doc.add_page_break()

    # 3. 画面一覧
    doc.add_heading('3. 画面一覧', 1)
    if analysis.pages:
        page_table = doc.add_table(rows=len(analysis.pages) + 1, cols=5)
        page_table.style = 'Table Grid'
        for j, h in enumerate(['No.', '画面名', '種別', 'ボタン数', '数式数']):
            page_table.rows[0].cells[j].text = h
        for i, p in enumerate(analysis.pages):
            page_table.rows[i + 1].cells[0].text = str(i + 1)
            page_table.rows[i + 1].cells[1].text = p.name
            page_table.rows[i + 1].cells[2].text = 'マスターページ' if p.page_type == 'masterPage' else 'ページ'
            page_table.rows[i + 1].cells[3].text = str(len(p.buttons))
            page_table.rows[i + 1].cells[4].text = str(len(p.formulas))

    # 4. ワークフロー
    if analysis.workflows:
        doc.add_page_break()
        doc.add_heading('4. ワークフロー定義', 1)
        for wf in analysis.workflows:
            doc.add_heading(f'ワークフロー: {wf.table_name}', 2)
            if wf.states:
                state_table = doc.add_table(rows=len(wf.states) + 1, cols=3)
                state_table.style = 'Table Grid'
                state_table.rows[0].cells[0].text = '状態名'
                state_table.rows[0].cells[1].text = '初期状態'
                state_table.rows[0].cells[2].text = '終了状態'
                for i, s in enumerate(wf.states):
                    state_table.rows[i + 1].cells[0].text = s.name
                    state_table.rows[i + 1].cells[1].text = '○' if s.is_initial else ''
                    state_table.rows[i + 1].cells[2].text = '○' if s.is_final else ''

    # 5. サーバーコマンド
    if analysis.server_commands:
        doc.add_page_break()
        doc.add_heading('5. サーバーコマンド', 1)
        cmd_list_table = doc.add_table(rows=len(analysis.server_commands) + 1, cols=4)
        cmd_list_table.style = 'Table Grid'
        for j, h in enumerate(['No.', 'コマンド名', 'フォルダ', 'パラメータ数']):
            cmd_list_table.rows[0].cells[j].text = h
        for i, c in enumerate(analysis.server_commands):
            cmd_list_table.rows[i + 1].cells[0].text = str(i + 1)
            cmd_list_table.rows[i + 1].cells[1].text = c.name
            cmd_list_table.rows[i + 1].cells[2].text = c.folder or '-'
            cmd_list_table.rows[i + 1].cells[3].text = str(len(c.parameters))

    file_path = os.path.join(output_dir, f'{analysis.project_name}_仕様書.docx')
    doc.save(file_path)
    return file_path


# =============================================================================
# Excel出力
# =============================================================================

def generate_excel_document(analysis: AnalysisResult, output_dir: str) -> str:
    """Excel形式で出力"""
    if not EXCEL_AVAILABLE:
        raise ImportError("openpyxlがインストールされていません。pip install openpyxl を実行してください。")

    os.makedirs(output_dir, exist_ok=True)
    wb = Workbook()

    # ヘッダースタイル
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    # サマリーシート
    ws_summary = wb.active
    ws_summary.title = 'サマリー'
    summary_data = [
        ['項目', '値'],
        ['プロジェクト名', analysis.project_name],
        ['テーブル数', analysis.summary.table_count],
        ['ページ数', analysis.summary.page_count],
        ['ワークフロー数', analysis.summary.workflow_count],
        ['サーバーコマンド数', analysis.summary.server_command_count],
        ['総カラム数', analysis.summary.total_columns],
        ['リレーション数', analysis.summary.total_relations],
        ['生成日', datetime.now().strftime('%Y-%m-%d %H:%M')],
        ['生成ツール', f'Forguncy Insight {VERSION_INFO}'],
    ]
    for row_idx, row in enumerate(summary_data, 1):
        for col_idx, value in enumerate(row, 1):
            cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx == 1:
                cell.fill = header_fill
                cell.font = header_font

    ws_summary.column_dimensions['A'].width = 20
    ws_summary.column_dimensions['B'].width = 40

    # テーブル一覧シート
    ws_tables = wb.create_sheet('テーブル一覧')
    table_headers = ['No.', 'テーブル名', 'フォルダ', 'カラム数', 'リレーション数']
    for col_idx, header in enumerate(table_headers, 1):
        cell = ws_tables.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    for row_idx, t in enumerate(analysis.tables, 2):
        values = [row_idx - 1, t.name, t.folder or '-', len(t.columns), len(t.relations)]
        for col_idx, value in enumerate(values, 1):
            cell = ws_tables.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    # カラム定義シート
    ws_columns = wb.create_sheet('カラム定義')
    col_headers = ['テーブル名', 'カラム名', 'データ型', '必須', 'ユニーク', 'デフォルト値']
    for col_idx, header in enumerate(col_headers, 1):
        cell = ws_columns.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    row_idx = 2
    for t in analysis.tables:
        for c in t.columns:
            values = [t.name, c.name, c.type, '○' if c.required else '', '○' if c.unique else '', c.default_value or '']
            for col_idx, value in enumerate(values, 1):
                cell = ws_columns.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
            row_idx += 1

    # ページ一覧シート
    ws_pages = wb.create_sheet('ページ一覧')
    page_headers = ['No.', 'ページ名', '種別', 'ボタン数', '数式数']
    for col_idx, header in enumerate(page_headers, 1):
        cell = ws_pages.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    for row_idx, p in enumerate(analysis.pages, 2):
        values = [row_idx - 1, p.name, 'マスターページ' if p.page_type == 'masterPage' else 'ページ', len(p.buttons), len(p.formulas)]
        for col_idx, value in enumerate(values, 1):
            cell = ws_pages.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    # サーバーコマンドシート
    ws_cmds = wb.create_sheet('サーバーコマンド')
    cmd_headers = ['No.', 'コマンド名', 'フォルダ', 'パラメータ数', '処理行数']
    for col_idx, header in enumerate(cmd_headers, 1):
        cell = ws_cmds.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    for row_idx, c in enumerate(analysis.server_commands, 2):
        values = [row_idx - 1, c.name, c.folder or '-', len(c.parameters), len(c.commands)]
        for col_idx, value in enumerate(values, 1):
            cell = ws_cmds.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    # ER図シート (テキスト形式)
    ws_er = wb.create_sheet('ER図(Mermaid)')
    ws_er.cell(row=1, column=1, value='以下をMermaid Live Editorに貼り付けてください')
    er_code = generate_er_mermaid(analysis.tables)
    ws_er.cell(row=3, column=1, value=er_code)

    file_path = os.path.join(output_dir, f'{analysis.project_name}_仕様書.xlsx')
    wb.save(file_path)
    return file_path


def generate_er_mermaid(tables: list) -> str:
    """ER図のMermaid記法を生成"""
    def sanitize(s):
        return re.sub(r'[^a-zA-Z0-9_\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF]', '_', s)

    lines = ['erDiagram']
    for table in tables:
        lines.append(f'  {sanitize(table.name)} {{')
        for col in table.columns[:10]:
            col_type = re.sub(r'[^a-z]', '', col.type.lower()) or 'string'
            pk = 'PK' if col.name.lower() == 'id' else ''
            lines.append(f'    {col_type} {sanitize(col.name)} {pk}'.strip())
        if len(table.columns) > 10:
            lines.append('    string more_columns "..."')
        lines.append('  }')

    for table in tables:
        for rel in table.relations:
            from_t = sanitize(table.name)
            to_t = sanitize(rel.target_table)
            rel_type = '}o--||' if 'Many' in rel.relation_type else '||--||'
            lines.append(f'  {from_t} {rel_type} {to_t} : "{rel.source_column}"')

    return '\n'.join(lines)


# =============================================================================
# GUI - モダンUI設定
# =============================================================================

# カラーパレット
COLORS = {
    "primary": "#3B82F6",        # ブルー
    "primary_hover": "#2563EB",  # ダークブルー
    "success": "#10B981",        # グリーン
    "warning": "#F59E0B",        # オレンジ
    "danger": "#EF4444",         # レッド
    "bg": "#F8FAFC",             # 背景
    "surface": "#FFFFFF",        # カード背景
    "text": "#1E293B",           # メインテキスト
    "text_secondary": "#64748B", # サブテキスト
    "text_muted": "#94A3B8",     # ミュートテキスト
    "border": "#E2E8F0",         # ボーダー
}

# フォント設定
FONT_FAMILY = "Yu Gothic UI"
FONTS = {
    "title": (FONT_FAMILY, 20, "bold"),
    "heading": (FONT_FAMILY, 14, "bold"),
    "body": (FONT_FAMILY, 11),
    "small": (FONT_FAMILY, 10),
}


class LicenseActivationDialog:
    """ライセンス認証ダイアログ（モダンUI）"""

    def __init__(self, parent: Tk, license_manager: LicenseManager):
        self.parent = parent
        self.license_manager = license_manager
        self.result = False

        self.dialog = Toplevel(parent)
        self.dialog.title("ライセンス認証")
        self.dialog.geometry("500x480")
        self.dialog.resizable(False, False)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.configure(bg=COLORS["bg"])

        # ダイアログを中央に配置
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() - 500) // 2
        y = (self.dialog.winfo_screenheight() - 480) // 2
        self.dialog.geometry(f"500x480+{x}+{y}")

        self.setup_ui()

        # 閉じるボタンの処理
        self.dialog.protocol("WM_DELETE_WINDOW", self.on_cancel)

    def setup_ui(self):
        # メインカード
        card = Frame(self.dialog, bg=COLORS["surface"], padx=40, pady=30)
        card.pack(fill='both', expand=True, padx=20, pady=20)

        # タイトル
        Label(card, text=PRODUCT_NAME, font=FONTS["title"],
              bg=COLORS["surface"], fg=COLORS["primary"]).pack(pady=(0, 5))
        Label(card, text="ライセンス認証", font=FONTS["heading"],
              bg=COLORS["surface"], fg=COLORS["text_secondary"]).pack(pady=(0, 25))

        # 説明
        desc_text = "製品をご利用いただくには、ライセンスキーの認証が必要です。"
        Label(card, text=desc_text, wraplength=400, justify='left',
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w', pady=(0, 20))

        # メールアドレス
        Label(card, text="メールアドレス:", anchor='w',
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(fill='x')
        self.email_entry = Entry(card, width=50, font=FONTS["body"], relief='solid', bd=1)
        self.email_entry.pack(fill='x', pady=(5, 15), ipady=5)

        # ライセンスキー
        Label(card, text="ライセンスキー:", anchor='w',
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(fill='x')
        self.key_entry = Entry(card, width=50, font=FONTS["body"], relief='solid', bd=1)
        self.key_entry.pack(fill='x', pady=(5, 8), ipady=5)
        Label(card, text="例: FGIN-STD-3101-XXXX-XXXX-XXXX",
              fg=COLORS["text_muted"], font=FONTS["small"], bg=COLORS["surface"]).pack(anchor='w')

        # エラーメッセージ
        self.error_label = Label(card, text="", fg=COLORS["danger"], wraplength=400,
                                  font=FONTS["small"], bg=COLORS["surface"])
        self.error_label.pack(pady=10)

        # ボタンフレーム
        btn_frame = Frame(card, bg=COLORS["surface"])
        btn_frame.pack(pady=15)

        Button(btn_frame, text="認証", command=self.on_activate,
               bg=COLORS["primary"], fg='white', font=FONTS["body"],
               padx=25, pady=8, relief='flat', cursor='hand2').pack(side='left', padx=5)
        Button(btn_frame, text="Free版で続行", command=self.on_continue_free,
               bg=COLORS["bg"], fg=COLORS["text"], font=FONTS["body"],
               padx=15, pady=8, relief='flat', cursor='hand2').pack(side='left', padx=5)

        # リンクフレーム
        link_frame = Frame(card, bg=COLORS["surface"])
        link_frame.pack(pady=10)

        trial_link = Label(link_frame, text="トライアル申請", fg=COLORS["primary"],
                           cursor='hand2', font=FONTS["small"], bg=COLORS["surface"])
        trial_link.pack(side='left', padx=10)
        trial_link.bind('<Button-1>', lambda e: webbrowser.open(TRIAL_URL))

        purchase_link = Label(link_frame, text="ライセンス購入", fg=COLORS["primary"],
                               cursor='hand2', font=FONTS["small"], bg=COLORS["surface"])
        purchase_link.pack(side='left', padx=10)
        purchase_link.bind('<Button-1>', lambda e: webbrowser.open(PURCHASE_URL))

        # 価格表示
        Label(card, text=f"Standard版: {PRICE_STANDARD}",
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text_secondary"]).pack(pady=5)

    def on_activate(self):
        email = self.email_entry.get().strip()
        key = self.key_entry.get().strip()

        if not email:
            self.error_label.config(text="メールアドレスを入力してください")
            return

        if not key:
            self.error_label.config(text="ライセンスキーを入力してください")
            return

        result = self.license_manager.activate(email, key)

        if result['is_valid']:
            self.result = True
            self.dialog.destroy()
            messagebox.showinfo("認証成功",
                f"ライセンスが正常に認証されました。\n\n"
                f"プラン: {self.license_manager.tier_name}\n"
                f"有効期限: {self.license_manager.expires_at.strftime('%Y年%m月%d日')}")
        else:
            self.error_label.config(text=result.get('error', '認証に失敗しました'))

    def on_continue_free(self):
        self.result = True
        self.dialog.destroy()

    def on_cancel(self):
        # Free版で続行として扱う
        self.on_continue_free()

    def show(self) -> bool:
        self.dialog.wait_window()
        return self.result


class ForguncyInsightApp:
    """モダンUI対応のメインアプリケーション"""

    def __init__(self, root: Tk):
        self.root = root
        self.root.title(f"Forguncy Insight {VERSION_INFO}")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        self.root.configure(bg=COLORS["bg"])

        # DPI対応
        try:
            import ctypes
            ctypes.windll.shcore.SetProcessDpiAwareness(1)
        except:
            pass

        self.license_manager = LicenseManager()
        self.file_path = StringVar()
        self.file_path2 = StringVar()  # 差分比較用
        self.output_dir = StringVar(value=str(Path.home() / "Documents"))

        self.setup_styles()
        self.setup_ui()

        # 起動時ライセンスチェック（UIセットアップ後）
        if not self.license_manager.is_activated:
            self.root.after(100, self._show_license_dialog)
        else:
            # 期限警告の表示
            self.show_expiry_warning()

    def setup_styles(self):
        """ttkスタイル設定"""
        style = ttk.Style()
        style.configure("TNotebook", background=COLORS["bg"])
        style.configure("TNotebook.Tab", font=FONTS["body"], padding=(15, 8))
        style.configure("TFrame", background=COLORS["surface"])
        style.configure("TLabel", background=COLORS["surface"], font=FONTS["body"])
        style.configure("TCheckbutton", background=COLORS["surface"], font=FONTS["body"])
        style.configure("TProgressbar", thickness=8)

    def _show_license_dialog(self):
        """ライセンスダイアログを表示"""
        dialog = LicenseActivationDialog(self.root, self.license_manager)
        dialog.show()
        # ダイアログ後にUI更新
        self.refresh_ui()
        self._update_license_badge()

    def _update_license_badge(self):
        """ライセンスバッジを更新"""
        tier = self.license_manager.tier
        badge_colors = {
            'TRIAL': (COLORS["warning"], "#FEF3C7"),
            'STD': (COLORS["primary"], "#DBEAFE"),
            'PRO': ("#8B5CF6", "#EDE9FE"),
        }
        if tier and tier in badge_colors:
            fg, bg = badge_colors[tier]
            self.license_badge.config(text=self.license_manager.tier_name, fg=fg, bg=bg)
        else:
            self.license_badge.config(text="Free", fg=COLORS["text_muted"], bg=COLORS["bg"])

    def show_expiry_warning(self):
        """期限警告を表示"""
        if self.license_manager.is_expiring_soon:
            warning_msg = self.license_manager.expiry_warning_message
            if warning_msg:
                messagebox.showwarning("ライセンス期限のお知らせ", warning_msg)

    def setup_ui(self):
        # ヘッダー
        header = Frame(self.root, bg=COLORS["surface"], height=70)
        header.pack(fill='x')
        header.pack_propagate(False)

        header_inner = Frame(header, bg=COLORS["surface"])
        header_inner.pack(fill='x', padx=20, pady=15)

        # タイトル
        title_frame = Frame(header_inner, bg=COLORS["surface"])
        title_frame.pack(side='left')

        Label(title_frame, text="◇ Forguncy Insight", font=FONTS["title"],
              bg=COLORS["surface"], fg=COLORS["primary"]).pack(side='left')

        Label(title_frame, text=f"  {VERSION_INFO}", font=FONTS["small"],
              bg=COLORS["surface"], fg=COLORS["text_muted"]).pack(side='left', padx=(10, 0))

        # ライセンスバッジ
        self.license_badge = Label(title_frame, text="Free", font=FONTS["small"],
                                    padx=10, pady=3, bg=COLORS["bg"], fg=COLORS["text_muted"])
        self.license_badge.pack(side='left', padx=(15, 0))
        self._update_license_badge()

        # ライセンスボタン
        license_btn = Button(header_inner, text="🔑 ライセンス", font=FONTS["small"],
                              bg=COLORS["bg"], fg=COLORS["text"], relief='flat',
                              padx=10, pady=5, cursor='hand2',
                              command=self._show_license_dialog)
        license_btn.pack(side='right')

        # Notebook (タブ)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=15, pady=(10, 15))

        # タブ1: 解析
        self.tab_analyze = Frame(self.notebook, bg=COLORS["surface"], padx=30, pady=25)
        self.notebook.add(self.tab_analyze, text='  解析  ')
        self.setup_analyze_tab()

        # タブ2: 差分比較
        self.tab_diff = Frame(self.notebook, bg=COLORS["surface"], padx=30, pady=25)
        self.notebook.add(self.tab_diff, text='  差分比較  ')
        self.setup_diff_tab()


    def setup_analyze_tab(self):
        # タイトル
        Label(self.tab_analyze, text="Forguncyプロジェクト解析・仕様書自動生成",
              font=FONTS["heading"], bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w', pady=(0, 20))

        # ドラッグ＆ドロップエリア
        self.drop_frame = Frame(self.tab_analyze, bg=COLORS["border"], padx=2, pady=2)
        self.drop_frame.pack(fill='x', pady=10)

        self.drop_area = Frame(self.drop_frame, bg="#F1F5F9", height=100)
        self.drop_area.pack(fill='both', expand=True)
        self.drop_area.pack_propagate(False)

        self.drop_label = Label(self.drop_area, text="📂 ここにファイルをドロップ\nまたはクリックして選択 (.fgcp)",
                                 font=FONTS["body"], bg="#F1F5F9", fg=COLORS["text_secondary"],
                                 cursor='hand2')
        self.drop_label.pack(expand=True)

        # ファイルパス表示
        self.file_label = Label(self.tab_analyze, textvariable=self.file_path,
                                 font=FONTS["small"], bg=COLORS["surface"], fg=COLORS["primary"],
                                 wraplength=600)
        self.file_label.pack(anchor='w', pady=(5, 0))

        # クリックイベント
        self.drop_area.bind('<Button-1>', lambda e: self.browse_file())
        self.drop_label.bind('<Button-1>', lambda e: self.browse_file())

        # ドラッグ＆ドロップ対応（tkinterdnd2があれば）
        self._setup_dnd()

        # 出力フォルダ
        output_frame = Frame(self.tab_analyze, bg=COLORS["surface"])
        output_frame.pack(fill='x', pady=10)
        Label(output_frame, text="出力フォルダ:",
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w')
        output_input = Frame(output_frame, bg=COLORS["surface"])
        output_input.pack(fill='x', pady=5)
        Entry(output_input, textvariable=self.output_dir,
              font=FONTS["body"], relief='solid', bd=1).pack(side='left', fill='x', expand=True, ipady=4)
        Button(output_input, text="変更...", command=self.browse_output,
               font=FONTS["body"], bg=COLORS["bg"], fg=COLORS["text"],
               relief='flat', padx=15, cursor='hand2').pack(side='right', padx=(10, 0))

        # 出力形式
        format_frame = Frame(self.tab_analyze, bg=COLORS["surface"])
        format_frame.pack(fill='x', pady=15)
        Label(format_frame, text="出力形式:",
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w')

        self.output_word = ttk.Checkbutton(format_frame, text="Word (.docx)")
        self.output_word.pack(anchor='w', pady=3)
        self.output_word.state(['selected'] if self.license_manager.limits.get('word_export') else ['disabled'])

        self.output_excel = ttk.Checkbutton(format_frame, text="Excel (.xlsx)")
        self.output_excel.pack(anchor='w')
        self.output_excel.state(['selected'] if self.license_manager.limits.get('excel_export') else ['disabled'])

        # プログレス
        progress_frame = Frame(self.tab_analyze, bg=COLORS["surface"])
        progress_frame.pack(fill='x', pady=20)
        self.progress = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress.pack(fill='x')
        self.status_label = Label(progress_frame, text="準備完了",
                                   font=FONTS["small"], bg=COLORS["surface"], fg=COLORS["text_muted"])
        self.status_label.pack(pady=(5, 0))

        # 解析ボタン
        self.analyze_btn = Button(self.tab_analyze, text="解析開始", command=self.start_analysis,
                                   font=FONTS["heading"], bg=COLORS["primary"], fg='white',
                                   padx=40, pady=12, relief='flat', cursor='hand2')
        self.analyze_btn.pack(pady=15)

        # Free版の制限表示
        if not self.license_manager.is_activated:
            limits = self.license_manager.limits
            limit_text = f"Free版制限: テーブル{int(limits['max_tables'])}件, ページ{int(limits['max_pages'])}件, サーバーコマンド{int(limits['max_server_commands'])}件"
            Label(self.tab_analyze, text=limit_text, fg=COLORS["text_muted"],
                  font=FONTS["small"], bg=COLORS["surface"]).pack()

    def setup_diff_tab(self):
        Label(self.tab_diff, text="プロジェクト差分比較", font=FONTS["heading"],
              bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w', pady=(0, 20))

        if not self.license_manager.limits.get('diff_compare'):
            Label(self.tab_diff, text="この機能はStandard版で利用できます",
                  fg=COLORS["danger"], font=FONTS["body"], bg=COLORS["surface"]).pack(pady=20)
            Button(self.tab_diff, text="ライセンスを購入",
                   command=lambda: webbrowser.open(PURCHASE_URL),
                   font=FONTS["body"], bg=COLORS["primary"], fg='white',
                   relief='flat', padx=20, cursor='hand2').pack()
            return

        # ファイル1（比較元）ドロップエリア
        Label(self.tab_diff, text="比較元ファイル (旧バージョン):",
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w')

        drop1_outer = Frame(self.tab_diff, bg=COLORS["border"], padx=2, pady=2)
        drop1_outer.pack(fill='x', pady=5)

        self.drop_area1 = Frame(drop1_outer, bg="#F1F5F9", height=70)
        self.drop_area1.pack(fill='both', expand=True)
        self.drop_area1.pack_propagate(False)

        self.drop_label1 = Label(self.drop_area1, text="📂 ファイルをドロップまたはクリック (.fgcp)",
                                  font=FONTS["body"], bg="#F1F5F9", fg=COLORS["text_secondary"],
                                  cursor='hand2')
        self.drop_label1.pack(expand=True)

        self.drop_area1.bind('<Button-1>', lambda e: self._browse_diff_file(1))
        self.drop_label1.bind('<Button-1>', lambda e: self._browse_diff_file(1))

        # ファイル2（比較先）ドロップエリア
        Label(self.tab_diff, text="比較先ファイル (新バージョン):",
              font=FONTS["body"], bg=COLORS["surface"], fg=COLORS["text"]).pack(anchor='w', pady=(15, 0))

        drop2_outer = Frame(self.tab_diff, bg=COLORS["border"], padx=2, pady=2)
        drop2_outer.pack(fill='x', pady=5)

        self.drop_area2 = Frame(drop2_outer, bg="#F1F5F9", height=70)
        self.drop_area2.pack(fill='both', expand=True)
        self.drop_area2.pack_propagate(False)

        self.drop_label2 = Label(self.drop_area2, text="📂 ファイルをドロップまたはクリック (.fgcp)",
                                  font=FONTS["body"], bg="#F1F5F9", fg=COLORS["text_secondary"],
                                  cursor='hand2')
        self.drop_label2.pack(expand=True)

        self.drop_area2.bind('<Button-1>', lambda e: self._browse_diff_file(2))
        self.drop_label2.bind('<Button-1>', lambda e: self._browse_diff_file(2))

        # 比較ボタン
        Button(self.tab_diff, text="差分を比較", command=self.compare_files,
               font=FONTS["heading"], bg=COLORS["success"], fg='white',
               padx=40, pady=12, relief='flat', cursor='hand2').pack(pady=30)


    def _browse_diff_file(self, file_num):
        """差分比較用ファイル選択"""
        path = filedialog.askopenfilename(title="Forguncyプロジェクトを選択", filetypes=[("Forguncy Project", "*.fgcp")])
        if path:
            if file_num == 1:
                self.file_path.set(path)
                self._update_diff_drop_area(1)
            else:
                self.file_path2.set(path)
                self._update_diff_drop_area(2)

    def _update_diff_drop_area(self, file_num):
        """差分用ドロップエリア表示更新"""
        if file_num == 1:
            path = self.file_path.get()
            label = self.drop_label1
            area = self.drop_area1
        else:
            path = self.file_path2.get()
            label = self.drop_label2
            area = self.drop_area2

        if path:
            filename = Path(path).name
            label.configure(text=f"✓ {filename}", bg="#ECFDF5", fg=COLORS["success"])
            area.configure(bg="#ECFDF5")

    def _setup_dnd(self):
        """ドラッグ＆ドロップを設定（tkinterdnd2が利用可能な場合）"""
        if not DND_AVAILABLE:
            return
        try:
            self.drop_area.drop_target_register(DND_FILES)
            self.drop_area.dnd_bind('<<Drop>>', self._on_drop)
            self.drop_area.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self.drop_area.dnd_bind('<<DragLeave>>', self._on_drag_leave)
        except Exception:
            pass

    def _on_drop(self, event):
        """ファイルドロップ時の処理"""
        path = event.data
        # Windowsでは{}で囲まれている場合がある
        if path.startswith('{') and path.endswith('}'):
            path = path[1:-1]
        if path.lower().endswith('.fgcp'):
            self.file_path.set(path)
            self._update_drop_area()
        self._on_drag_leave(None)

    def _on_drag_enter(self, event):
        """ドラッグ中の表示"""
        self.drop_area.configure(bg="#DBEAFE")
        self.drop_label.configure(bg="#DBEAFE", fg=COLORS["primary"])

    def _on_drag_leave(self, event):
        """ドラッグ離脱時の表示"""
        if self.file_path.get():
            self.drop_area.configure(bg="#ECFDF5")
            self.drop_label.configure(bg="#ECFDF5", fg=COLORS["success"])
        else:
            self.drop_area.configure(bg="#F1F5F9")
            self.drop_label.configure(bg="#F1F5F9", fg=COLORS["text_secondary"])

    def _update_drop_area(self):
        """ファイル選択後のドロップエリア表示更新"""
        if self.file_path.get():
            filename = Path(self.file_path.get()).name
            self.drop_label.configure(text=f"✓ {filename}\n（クリックで変更）",
                                       bg="#ECFDF5", fg=COLORS["success"])
            self.drop_area.configure(bg="#ECFDF5")

    def browse_file(self):
        path = filedialog.askopenfilename(title="Forguncyプロジェクトを選択", filetypes=[("Forguncy Project", "*.fgcp")])
        if path:
            self.file_path.set(path)
            self._update_drop_area()

    def browse_file2(self):
        path = filedialog.askopenfilename(title="比較先プロジェクトを選択", filetypes=[("Forguncy Project", "*.fgcp")])
        if path:
            self.file_path2.set(path)

    def browse_output(self):
        path = filedialog.askdirectory(title="出力フォルダを選択", initialdir=self.output_dir.get())
        if path:
            self.output_dir.set(path)

    def update_progress(self, pct: int, msg: str):
        self.progress['value'] = pct
        self.status_label.config(text=msg)
        self.root.update_idletasks()

    def start_analysis(self):
        if not self.file_path.get():
            messagebox.showerror("エラー", "プロジェクトファイルを選択してください")
            return

        self.analyze_btn.config(state='disabled')
        self.progress['value'] = 0

        try:
            self.update_progress(10, "解析を開始しています...")
            analysis = analyze_project(self.file_path.get(), self.update_progress, self.license_manager.limits)

            generated_files = []

            # Word出力
            if self.license_manager.limits.get('word_export'):
                self.update_progress(70, "Word仕様書を生成しています...")
                word_path = generate_spec_document(analysis, self.output_dir.get())
                generated_files.append(word_path)

            # Excel出力
            if self.license_manager.limits.get('excel_export') and EXCEL_AVAILABLE:
                self.update_progress(85, "Excel仕様書を生成しています...")
                excel_path = generate_excel_document(analysis, self.output_dir.get())
                generated_files.append(excel_path)

            self.update_progress(100, "完了しました!")

            msg = f"解析が完了しました。\n\nテーブル: {analysis.summary.table_count}件\nページ: {analysis.summary.page_count}件"
            if generated_files:
                msg += f"\n\n生成ファイル:\n" + "\n".join(generated_files)
            else:
                msg += "\n\n※ Word/Excel出力にはStandard版が必要です"

            messagebox.showinfo("完了", msg)

            if generated_files and os.name == 'nt':
                os.startfile(self.output_dir.get())

        except Exception as e:
            messagebox.showerror("エラー", f"解析中にエラーが発生しました:\n{str(e)}")
            self.update_progress(0, "")
        finally:
            self.analyze_btn.config(state='normal')

    def compare_files(self):
        if not self.file_path.get() or not self.file_path2.get():
            messagebox.showerror("エラー", "2つのファイルを選択してください")
            return

        try:
            old_analysis = analyze_project(self.file_path.get(), limits=self.license_manager.limits)
            new_analysis = analyze_project(self.file_path2.get(), limits=self.license_manager.limits)
            diff = compare_projects(old_analysis, new_analysis)

            # 結果表示ウィンドウ
            diff_window = Toplevel(self.root)
            diff_window.title("差分比較結果")
            diff_window.geometry("600x500")

            text = Text(diff_window, wrap=WORD, padx=10, pady=10)
            scrollbar = Scrollbar(diff_window, command=text.yview)
            text.configure(yscrollcommand=scrollbar.set)
            scrollbar.pack(side=RIGHT, fill=Y)
            text.pack(fill=BOTH, expand=True)

            text.insert(END, f"=== 差分比較結果 ===\n\n")
            text.insert(END, f"比較元: {old_analysis.project_name}\n")
            text.insert(END, f"比較先: {new_analysis.project_name}\n\n")

            text.insert(END, f"--- テーブル ---\n")
            text.insert(END, f"追加: {len(diff.added_tables)}件\n")
            for t in diff.added_tables:
                text.insert(END, f"  + {t.name}\n")
            text.insert(END, f"削除: {len(diff.removed_tables)}件\n")
            for t in diff.removed_tables:
                text.insert(END, f"  - {t.name}\n")
            text.insert(END, f"変更: {len(diff.modified_tables)}件\n")
            for m in diff.modified_tables:
                text.insert(END, f"  * {m['name']}\n")

            text.insert(END, f"\n--- ページ ---\n")
            text.insert(END, f"追加: {len(diff.added_pages)}件\n")
            for p in diff.added_pages:
                text.insert(END, f"  + {p.name}\n")
            text.insert(END, f"削除: {len(diff.removed_pages)}件\n")
            for p in diff.removed_pages:
                text.insert(END, f"  - {p.name}\n")

            text.insert(END, f"\n--- サーバーコマンド ---\n")
            text.insert(END, f"追加: {len(diff.added_server_commands)}件\n")
            for c in diff.added_server_commands:
                text.insert(END, f"  + {c.name}\n")
            text.insert(END, f"削除: {len(diff.removed_server_commands)}件\n")
            for c in diff.removed_server_commands:
                text.insert(END, f"  - {c.name}\n")
            text.insert(END, f"変更: {len(diff.modified_server_commands)}件\n")
            for m in diff.modified_server_commands:
                text.insert(END, f"  * {m['name']}\n")

            text.config(state='disabled')

        except Exception as e:
            messagebox.showerror("エラー", f"比較中にエラーが発生しました:\n{str(e)}")


    def refresh_ui(self):
        """ライセンス状態に応じてUIを更新"""
        # 出力チェックボックスの状態を更新
        if self.license_manager.limits.get('word_export'):
            self.output_word.state(['!disabled', 'selected'])
        else:
            self.output_word.state(['disabled', '!selected'])

        if self.license_manager.limits.get('excel_export'):
            self.output_excel.state(['!disabled', 'selected'])
        else:
            self.output_excel.state(['disabled', '!selected'])

        # 差分比較タブを再構築
        for widget in self.tab_diff.winfo_children():
            widget.destroy()
        self.setup_diff_tab()

        # Free版制限表示の更新
        self._update_license_badge()


def main():
    # ドラッグ＆ドロップを有効にするためTkinterDnD.Tk()を使用
    if DND_AVAILABLE:
        root = TkinterDnD.Tk()
    else:
        root = Tk()
    app = ForguncyInsightApp(root)
    root.mainloop()


if __name__ == '__main__':
    main()
