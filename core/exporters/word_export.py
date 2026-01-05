#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word仕様書エクスポートモジュール

詳細仕様書をWord形式で出力する。
"""

import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

from core.models import AnalysisResult, CommandInfo


# =============================================================================
# バージョン情報（メインモジュールから参照）
# =============================================================================
APP_VERSION = "1.1.0"
SUPPORTED_FORGUNCY_VERSIONS = ["9.x"]
VERSION_INFO = f"v{APP_VERSION} (Forguncy {', '.join(SUPPORTED_FORGUNCY_VERSIONS)} 対応)"


# =============================================================================
# ヘルパー関数
# =============================================================================
def _classify_command_type(commands: list) -> str:
    """コマンドリストからコマンド種別を判定"""
    if not commands:
        return '-'
    types = set()
    for cmd in commands:
        # CommandInfoオブジェクトの場合はdescriptionを使用
        if isinstance(cmd, CommandInfo):
            cmd_lower = cmd.description.lower() if cmd.description else ''
        elif isinstance(cmd, str):
            cmd_lower = cmd.lower()
        else:
            cmd_lower = ''

        if 'navigate' in cmd_lower or 'page:' in cmd_lower:
            types.add('画面遷移')
        elif 'update' in cmd_lower or 'insert' in cmd_lower or 'delete' in cmd_lower:
            types.add('データ操作')
        elif 'query' in cmd_lower or 'select' in cmd_lower:
            types.add('データ取得')
        elif 'call' in cmd_lower or 'server' in cmd_lower:
            types.add('サーバー呼出')
        elif 'message' in cmd_lower or 'alert' in cmd_lower:
            types.add('メッセージ')
        elif 'export' in cmd_lower or 'import' in cmd_lower:
            types.add('入出力')
    return ', '.join(types) if types else 'その他'


def _set_table_header_style(table, header_row_idx=0):
    """テーブルヘッダー行のスタイル設定"""
    for cell in table.rows[header_row_idx].cells:
        cell._element.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '4472C4')
        cell._element.tcPr.append(shading)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)


# =============================================================================
# Word仕様書生成
# =============================================================================
def generate_spec_document(analysis: AnalysisResult, output_dir: str) -> str:
    """詳細仕様書ドキュメントを生成"""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # スタイル設定
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    style.font.size = Pt(11)

    # ================== 表紙 ==================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading(analysis.project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('詳細システム仕様書')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
    doc.add_paragraph()
    doc.add_paragraph(f'作成日: {datetime.now().strftime("%Y/%m/%d")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Forguncy Insight {VERSION_INFO} 自動生成').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ================== 目次 ==================
    doc.add_heading('目次', 1)
    toc_items = [
        '1. システム概要',
        '   1.1 システム構成',
        '2. 画面一覧',
    ]
    # 画面フォルダをTOCに追加
    page_folders = sorted(set(p.folder for p in analysis.pages if p.folder))
    for i, folder in enumerate(page_folders[:10], 1):
        toc_items.append(f'   2.{i} {folder}')
    toc_items.append('3. テーブル定義')
    table_folders = sorted(set(t.folder for t in analysis.tables if t.folder))
    for i, folder in enumerate(table_folders[:10], 1):
        toc_items.append(f'   3.{i} {folder}')
    if analysis.workflows:
        toc_items.append('4. ワークフロー定義')
    if analysis.server_commands:
        toc_items.append('5. サーバーコマンド')
    # ボタン・コマンドがある画面をカウント
    pages_with_buttons = [p for p in analysis.pages if p.buttons]
    if pages_with_buttons:
        toc_items.append('6. ボタン・コマンド詳細')

    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ================== 1. システム概要 ==================
    doc.add_heading('1. システム概要', 1)
    doc.add_paragraph('本ドキュメントは Forguncy プロジェクトの詳細システム仕様書です。')

    doc.add_heading('1.1 システム構成', 2)
    summary_table = doc.add_table(rows=9, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('項目', '値'),
        ('プロジェクト名', analysis.project_name),
        ('画面数', f'{analysis.summary.page_count} 画面'),
        ('テーブル数', f'{analysis.summary.table_count} テーブル'),
        ('サーバーコマンド数', f'{analysis.summary.server_command_count} 件'),
        ('ワークフロー数', f'{analysis.summary.workflow_count} 件'),
        ('総カラム数', f'{analysis.summary.total_columns} 件'),
        ('リレーション数', f'{analysis.summary.total_relations} 件'),
        ('対応バージョン', ', '.join(SUPPORTED_FORGUNCY_VERSIONS)),
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    _set_table_header_style(summary_table, 0)

    doc.add_page_break()

    # ================== 2. 画面一覧 ==================
    doc.add_heading('2. 画面一覧', 1)

    # フォルダ別にグループ化
    pages_by_folder = {}
    for p in analysis.pages:
        folder = p.folder or '(ルート)'
        if folder not in pages_by_folder:
            pages_by_folder[folder] = []
        pages_by_folder[folder].append(p)

    section_num = 1
    for folder in sorted(pages_by_folder.keys()):
        folder_pages = pages_by_folder[folder]
        doc.add_heading(f'2.{section_num} {folder}', 2)

        page_table = doc.add_table(rows=len(folder_pages) + 1, cols=4)
        page_table.style = 'Table Grid'
        headers = ['No.', '画面名', 'ボタン数', '数式数']
        for j, h in enumerate(headers):
            page_table.rows[0].cells[j].text = h
        _set_table_header_style(page_table, 0)

        for i, p in enumerate(folder_pages):
            page_table.rows[i + 1].cells[0].text = str(i + 1)
            page_table.rows[i + 1].cells[1].text = p.name
            page_table.rows[i + 1].cells[2].text = str(len(p.buttons))
            page_table.rows[i + 1].cells[3].text = str(len(p.formulas))

        section_num += 1

    doc.add_page_break()

    # ================== 3. テーブル定義 ==================
    doc.add_heading('3. テーブル定義', 1)

    # フォルダ別にグループ化
    tables_by_folder = {}
    for t in analysis.tables:
        folder = t.folder or '(ルート)'
        if folder not in tables_by_folder:
            tables_by_folder[folder] = []
        tables_by_folder[folder].append(t)

    table_section_num = 1
    table_detail_num = 1
    for folder in sorted(tables_by_folder.keys()):
        folder_tables = tables_by_folder[folder]
        doc.add_heading(f'3.{table_section_num} {folder}', 2)

        for table in folder_tables:
            doc.add_heading(f'3.{table_section_num}.{table_detail_num} {table.name}', 3)

            # テーブル概要
            if table.primary_key:
                doc.add_paragraph(f'主キー: {", ".join(table.primary_key)}')
            if table.relations:
                rel_text = ', '.join([f'{r.target_table}({r.source_column}→{r.target_column})' for r in table.relations[:3]])
                if len(table.relations) > 3:
                    rel_text += f' 他{len(table.relations) - 3}件'
                doc.add_paragraph(f'リレーション: {rel_text}')

            # カラム定義テーブル
            if table.columns:
                col_table = doc.add_table(rows=len(table.columns) + 1, cols=6)
                col_table.style = 'Table Grid'
                headers = ['No.', 'カラム名', 'データ型', '主キー', '必須', 'デフォルト値']
                for j, h in enumerate(headers):
                    col_table.rows[0].cells[j].text = h
                _set_table_header_style(col_table, 0)

                for i, c in enumerate(table.columns):
                    col_table.rows[i + 1].cells[0].text = str(i + 1)
                    col_table.rows[i + 1].cells[1].text = c.name
                    col_table.rows[i + 1].cells[2].text = c.type
                    col_table.rows[i + 1].cells[3].text = '○' if c.name in table.primary_key else ''
                    col_table.rows[i + 1].cells[4].text = '○' if c.required else ''
                    col_table.rows[i + 1].cells[5].text = c.default_value or ''

            table_detail_num += 1

        table_section_num += 1
        table_detail_num = 1

    doc.add_page_break()

    # ================== 4. ワークフロー定義 ==================
    if analysis.workflows:
        doc.add_heading('4. ワークフロー定義', 1)
        doc.add_paragraph('本システムで定義されているワークフローの詳細を以下に示します。')

        for idx, wf in enumerate(analysis.workflows, 1):
            doc.add_heading(f'4.{idx} {wf.table_name}', 2)

            if wf.states:
                doc.add_paragraph('■ 状態一覧')
                state_table = doc.add_table(rows=len(wf.states) + 1, cols=4)
                state_table.style = 'Table Grid'
                headers = ['No.', '状態名', '初期状態', '終了状態']
                for j, h in enumerate(headers):
                    state_table.rows[0].cells[j].text = h
                _set_table_header_style(state_table, 0)

                for i, s in enumerate(wf.states):
                    state_table.rows[i + 1].cells[0].text = str(i + 1)
                    state_table.rows[i + 1].cells[1].text = s.name
                    state_table.rows[i + 1].cells[2].text = '○' if s.is_initial else ''
                    state_table.rows[i + 1].cells[3].text = '○' if s.is_final else ''

            if wf.transitions:
                doc.add_paragraph()
                doc.add_paragraph('■ 遷移一覧')
                trans_table = doc.add_table(rows=len(wf.transitions) + 1, cols=4)
                trans_table.style = 'Table Grid'
                headers = ['No.', '遷移名', '遷移元', '遷移先']
                for j, h in enumerate(headers):
                    trans_table.rows[0].cells[j].text = h
                _set_table_header_style(trans_table, 0)

                for i, t in enumerate(wf.transitions):
                    trans_table.rows[i + 1].cells[0].text = str(i + 1)
                    trans_table.rows[i + 1].cells[1].text = t.action
                    trans_table.rows[i + 1].cells[2].text = t.from_state
                    trans_table.rows[i + 1].cells[3].text = t.to_state

        doc.add_page_break()
    else:
        doc.add_heading('4. ワークフロー定義', 1)
        doc.add_paragraph('ワークフローは定義されていません。')
        doc.add_page_break()

    # ================== 5. サーバーコマンド ==================
    if analysis.server_commands:
        doc.add_heading('5. サーバーコマンド', 1)
        doc.add_paragraph('サーバー側で実行されるコマンドの一覧と詳細を以下に示します。')

        # 一覧テーブル
        cmd_table = doc.add_table(rows=len(analysis.server_commands) + 1, cols=4)
        cmd_table.style = 'Table Grid'
        headers = ['No.', 'コマンド名', 'フォルダ', 'パラメータ数']
        for j, h in enumerate(headers):
            cmd_table.rows[0].cells[j].text = h
        _set_table_header_style(cmd_table, 0)

        for i, c in enumerate(analysis.server_commands):
            cmd_table.rows[i + 1].cells[0].text = str(i + 1)
            cmd_table.rows[i + 1].cells[1].text = c.name
            cmd_table.rows[i + 1].cells[2].text = c.folder or '-'
            cmd_table.rows[i + 1].cells[3].text = str(len(c.parameters))

        # 各コマンドの詳細
        for idx, cmd in enumerate(analysis.server_commands, 1):
            doc.add_heading(f'5.{idx} {cmd.name}', 2)

            if cmd.parameters:
                doc.add_paragraph('■ パラメータ')
                param_table = doc.add_table(rows=len(cmd.parameters) + 1, cols=4)
                param_table.style = 'Table Grid'
                headers = ['No.', 'パラメータ名', 'データ型', '必須']
                for j, h in enumerate(headers):
                    param_table.rows[0].cells[j].text = h
                _set_table_header_style(param_table, 0)

                for i, p in enumerate(cmd.parameters):
                    param_table.rows[i + 1].cells[0].text = str(i + 1)
                    param_table.rows[i + 1].cells[1].text = p.name
                    param_table.rows[i + 1].cells[2].text = p.type or '-'
                    param_table.rows[i + 1].cells[3].text = '○' if p.required else ''

            if cmd.commands:
                doc.add_paragraph()
                doc.add_paragraph('■ 処理内容')
                for c in cmd.commands[:10]:
                    # CommandInfoオブジェクトの場合
                    if isinstance(c, CommandInfo):
                        cmd_text = c.description
                    else:
                        cmd_text = str(c)
                    doc.add_paragraph(f'  • {cmd_text[:100]}' if len(cmd_text) > 100 else f'  • {cmd_text}')
                if len(cmd.commands) > 10:
                    doc.add_paragraph(f'  ... 他 {len(cmd.commands) - 10} 件')

        doc.add_page_break()

    # ================== 6. ボタン・コマンド詳細 ==================
    pages_with_buttons = [p for p in analysis.pages if p.buttons]
    if pages_with_buttons:
        doc.add_heading('6. ボタン・コマンド詳細', 1)
        doc.add_paragraph('各画面のボタンに設定されているコマンドの詳細を以下に示します。')

        for idx, page in enumerate(pages_with_buttons, 1):
            doc.add_heading(f'6.{idx} {page.name}', 2)

            btn_table = doc.add_table(rows=len(page.buttons) + 1, cols=3)
            btn_table.style = 'Table Grid'
            headers = ['ボタン名', 'コマンド種別', '詳細']
            for j, h in enumerate(headers):
                btn_table.rows[0].cells[j].text = h
            _set_table_header_style(btn_table, 0)

            for i, btn in enumerate(page.buttons):
                btn_table.rows[i + 1].cells[0].text = btn.name or '(名称なし)'
                btn_table.rows[i + 1].cells[1].text = _classify_command_type(btn.commands)
                # 詳細は最初の2つのコマンドのdescriptionを表示
                detail = ', '.join(cmd.description for cmd in btn.commands[:2]) if btn.commands else '-'
                if len(detail) > 60:
                    detail = detail[:57] + '...'
                btn_table.rows[i + 1].cells[2].text = detail

            # ボタンが多い場合の補足
            if len(page.buttons) > 20:
                doc.add_paragraph(f'※ 画面内に {len(page.buttons)} 個のボタンがあります')

    # 保存
    file_path = os.path.join(output_dir, f'{analysis.project_name}_詳細仕様書.docx')
    doc.save(file_path)
    return file_path
